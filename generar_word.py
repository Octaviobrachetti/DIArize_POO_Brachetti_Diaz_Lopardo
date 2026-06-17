# -*- coding: utf-8 -*-
"""Genera el documento Word con la explicacion tecnica de POO del proyecto DIArize."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AZUL = RGBColor(0x06, 0x4E, 0x72)
CYAN = RGBColor(0x0E, 0x74, 0x90)
GRIS = RGBColor(0x33, 0x41, 0x55)
GRIS_CLARO = RGBColor(0xF1, 0xF5, 0xF9)
NEGRO_CODE = RGBColor(0x1E, 0x29, 0x3B)

doc = Document()

# --- Estilos base ---
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

def _shade(par, color_hex):
    pPr = par._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    pPr.append(shd)

def titulo_portada(texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(texto)
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = AZUL
    return p

def subtitulo_portada(texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(texto)
    r.font.size = Pt(14)
    r.font.color.rgb = GRIS
    return p

def h1(texto):
    p = doc.add_heading(level=1)
    r = p.add_run(texto)
    r.font.color.rgb = AZUL
    r.font.size = Pt(18)
    r.font.bold = True
    return p

def h2(texto):
    p = doc.add_heading(level=2)
    r = p.add_run(texto)
    r.font.color.rgb = CYAN
    r.font.size = Pt(14)
    r.font.bold = True
    return p

def parrafo(texto, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.bold = bold
    return p

def etiqueta(label, texto):
    """Parrafo con una etiqueta en negrita azul seguida del texto."""
    p = doc.add_paragraph()
    r = p.add_run(label + " ")
    r.font.bold = True
    r.font.color.rgb = CYAN
    r2 = p.add_run(texto)
    return p

def bullet(texto):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(texto)
    return p

def codigo(texto):
    """Bloque de codigo con fondo gris y fuente monoespaciada."""
    p = doc.add_paragraph()
    _shade(p, "1E293B")
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    for i, linea in enumerate(texto.split("\n")):
        r = p.add_run(("\n" if i else "") + linea)
        r.font.name = "Consolas"
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
    return p

def separador():
    p = doc.add_paragraph()
    r = p.add_run("─" * 60)
    r.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============================================================ PORTADA
doc.add_paragraph()
doc.add_paragraph()
titulo_portada("DIArize")
subtitulo_portada("Transcriptor y diarizador de audio con IA")
doc.add_paragraph()
subtitulo_portada("Conceptos de Programacion Orientada a Objetos aplicados")
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Brachetti  •  Diaz  •  Lopardo")
r.font.size = Pt(13)
r.font.bold = True
r.font.color.rgb = GRIS
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Trabajo Final - Programacion Orientada a Objetos")
r.font.size = Pt(11)
r.font.color.rgb = GRIS

doc.add_page_break()

# ============================================================ INTRO / ARQUITECTURA
h1("1. Vision general del proyecto")
parrafo(
    "DIArize es una aplicacion de escritorio que transcribe audio (desde un archivo o en vivo "
    "por microfono), separa quien dijo cada cosa (diarizacion) y permite procesar el texto con "
    "inteligencia artificial: limpiarlo, resumirlo, analizarlo, traducirlo y hacerle preguntas."
)
parrafo("El proyecto esta dividido en TRES capas bien separadas:", bold=True)

tabla = doc.add_table(rows=1, cols=3)
tabla.style = "Light Grid Accent 1"
hdr = tabla.rows[0].cells
hdr[0].paragraphs[0].add_run("Capa").bold = True
hdr[1].paragraphs[0].add_run("Tecnologia").bold = True
hdr[2].paragraphs[0].add_run("Responsabilidad").bold = True
filas = [
    ("Nucleo (nucleo/)", "C++ puro (sin Qt)",
     "Modelo de datos y contratos: Hablante, Segmento, Transcripcion, Repositorio, "
     "interfaces ITranscriptor e IObservador."),
    ("Red + GUI (app/)", "C++ con Qt",
     "Clientes HTTP/WebSocket que hablan con el servidor, y la ventana grafica."),
    ("Servidor (Transcriptor/)", "Python + Flask",
     "Recibe pedidos del cliente y habla con las APIs externas (AssemblyAI y Gemini/OpenAI)."),
]
for a, b, c in filas:
    row = tabla.add_row().cells
    row[0].paragraphs[0].add_run(a)
    row[1].paragraphs[0].add_run(b)
    row[2].paragraphs[0].add_run(c)

doc.add_paragraph()
etiqueta("Por que separar el nucleo de Qt:",
    "el nucleo no depende de ninguna libreria grafica, asi que es portable y se puede probar "
    "sin levantar la interfaz. Esto se llama BAJO ACOPLAMIENTO: cada parte depende lo menos "
    "posible de las demas.")

doc.add_page_break()

# ============================================================ CONCEPTOS POO
h1("2. Conceptos de POO aplicados")
parrafo(
    "Para cada concepto explicamos QUE ES, DONDE lo usamos (con codigo real), COMO funciona, "
    "POR QUE lo elegimos y POR QUE no usamos otra alternativa."
)

# ---------- ENCAPSULAMIENTO ----------
h2("2.1 Encapsulamiento")
etiqueta("Que es:",
    "es ocultar los datos internos de una clase (atributos) y permitir el acceso solo a traves "
    "de metodos publicos controlados. La clase se vuelve la unica dueña de su estado.")
etiqueta("Donde lo usamos:",
    "en todas las clases del nucleo. El ejemplo mas claro es la clase Hablante "
    "(nucleo/include/nucleo/Hablante.h):")
codigo(
"class Hablante {\n"
"public:\n"
"    Hablante(const std::string& id, const std::string& nombre = \"\");\n"
"    const std::string& getId()     const;\n"
"    const std::string& getNombre() const;\n"
"    void               setNombre(const std::string& nombre);\n"
"    double getTiempoTotal() const;\n"
"    void   agregarTiempo(double seg);\n"
"private:\n"
"    std::string _id;\n"
"    std::string _nombre;\n"
"    double      _tiempoTotal { 0.0 };\n"
"};"
)
etiqueta("Como funciona:",
    "los atributos (_id, _nombre, _tiempoTotal) estan en la seccion 'private', por lo que NADIE "
    "de afuera puede tocarlos directamente. Para leerlos se usan getters (getNombre) y para "
    "modificarlos, setters o metodos especificos (setNombre, agregarTiempo).")
etiqueta("Por que lo usamos:",
    "para proteger la integridad de los datos. Dos ejemplos concretos: (1) _tiempoTotal solo se "
    "puede aumentar con agregarTiempo(seg), nunca asignar un valor invalido como un numero "
    "negativo; (2) _id NO tiene setter, asi que es inmutable: la identidad de un hablante no "
    "deberia cambiar despues de crearlo.")
etiqueta("Por que NO la otra opcion:",
    "si hubieramos dejado los atributos publicos, cualquier parte del programa podria corromper "
    "el estado del objeto (por ejemplo poner un tiempo negativo) y seria imposible rastrear "
    "donde se rompio. El encapsulamiento centraliza el control en la propia clase.")

# ---------- ABSTRACCION ----------
h2("2.2 Abstraccion (clases abstractas / interfaces)")
etiqueta("Que es:",
    "definir QUE hace algo sin decir COMO lo hace. Una clase abstracta es un 'contrato': declara "
    "metodos que las clases hijas estan obligadas a implementar, pero no da la implementacion.")
etiqueta("Donde lo usamos:",
    "en las dos interfaces del proyecto: ITranscriptor e IObservador. El prefijo 'I' significa "
    "'Interfaz' por convencion. Veamos ITranscriptor (nucleo/include/nucleo/ITranscriptor.h):")
codigo(
"class ITranscriptor {\n"
"public:\n"
"    virtual ~ITranscriptor() = default;\n"
"    // Funcion virtual PURA: obliga a las subclases a implementarla\n"
"    virtual Transcripcion transcribir(const std::string& rutaAudio,\n"
"                                      const std::string& idioma = \"es\",\n"
"                                      int numHablantes = 0) = 0;\n"
"    // Funcion virtual con cuerpo: las subclases PUEDEN sobreescribirla\n"
"    virtual bool estaDisponible() const { return true; }\n"
"};"
)
etiqueta("Como funciona:",
    "ITranscriptor dice 'un transcriptor sabe transcribir', pero no contiene la logica de COMO. "
    "Eso lo decide quien la implemente: ClienteTranscriptor lo hace enviando el audio por HTTP, "
    "pero podria existir un TranscriptorLocal que use Whisper offline. La interfaz expone lo "
    "esencial y oculta el detalle.")
etiqueta("Por que lo usamos:",
    "para separar el concepto de su implementacion. El resto del programa trabaja con la idea "
    "abstracta 'transcriptor' sin que le importe que hay por debajo. Eso hace el codigo flexible "
    "y preparado para cambios futuros sin tener que reescribir todo.")
etiqueta("Por que NO la otra opcion:",
    "si la GUI dependiera directamente de una clase concreta (por ejemplo llamar siempre a "
    "AssemblyAI), cambiar de proveedor o agregar un transcriptor de prueba obligaria a modificar "
    "la GUI. Con la interfaz, la GUI no se entera del cambio.")

# ---------- HERENCIA SIMPLE ----------
h2("2.3 Herencia simple")
etiqueta("Que es:",
    "una clase (hija) hereda atributos y metodos de otra clase (padre o base), reutilizando su "
    "comportamiento y agregando o especializando lo propio. Es la relacion 'es-un'.")
etiqueta("Donde lo usamos:",
    "ClienteTranscriptor hereda de ClienteHTTP. ClienteHTTP (app/include/red/ClienteHTTP.h) es "
    "la clase base que encapsula toda la mecanica de red:")
codigo(
"class ClienteHTTP {\n"
"public:\n"
"    explicit ClienteHTTP(const QString& urlBase);\n"
"    virtual ~ClienteHTTP();\n"
"    QByteArray get (const QString& endpoint) const;\n"
"    QByteArray post(const QString& endpoint, QHttpMultiPart* multipart) const;\n"
"    QByteArray postJson(const QString& endpoint, const QByteArray& json) const;\n"
"protected:\n"
"    virtual QNetworkRequest construirRequest(const QUrl& url) const;\n"
"    QString _urlBase;\n"
"    mutable QNetworkAccessManager _manager;\n"
"};"
)
etiqueta("Como funciona:",
    "ClienteTranscriptor recibe automaticamente los metodos get(), post() y postJson() de "
    "ClienteHTTP. No los reescribe: los reutiliza. Solo le suma lo suyo (transcribir y las "
    "operaciones de IA).")
etiqueta("Por que lo usamos:",
    "para reutilizar codigo sin copiarlo. Toda la logica de red (manejar QNetworkAccessManager, "
    "esperar la respuesta con un QEventLoop, leer los datos) esta escrita UNA sola vez en "
    "ClienteHTTP. Un ClienteTranscriptor 'es un' cliente HTTP especializado.")
etiqueta("Por que NO la otra opcion:",
    "sin herencia tendriamos que duplicar todo el codigo de red dentro de ClienteTranscriptor. "
    "Si manana hay un bug en como se envia un POST, habria que arreglarlo en cada copia. Con "
    "herencia se arregla en un solo lugar.")

# ---------- HERENCIA MULTIPLE ----------
h2("2.4 Herencia multiple (el caso mas importante)")
etiqueta("Que es:",
    "una clase hereda de DOS o mas clases base al mismo tiempo, combinando lo de todas.")
etiqueta("Donde lo usamos:",
    "en dos lugares clave del proyecto:")
codigo(
"// app/include/red/ClienteTranscriptor.h\n"
"class ClienteTranscriptor : public ClienteHTTP,                 // implementacion de red\n"
"                            public DIArize::Core::ITranscriptor // contrato del dominio\n"
"\n"
"// app/include/gui/VentanaPrincipal.h\n"
"class VentanaPrincipal : public QMainWindow,                    // es una ventana\n"
"                         public DIArize::Core::IObservador      // y un observador"
)
etiqueta("Como funciona:",
    "ClienteTranscriptor obtiene a la vez los metodos HTTP de ClienteHTTP Y el contrato de "
    "ITranscriptor. Asi puede enviar requests (gracias a la base de red) y ademas ser tratado "
    "como 'un transcriptor cualquiera' (gracias a la interfaz).")
etiqueta("Por que lo usamos (clave para defender):",
    "porque combinamos dos cosas de naturaleza DISTINTA que no tienen relacion entre si. De "
    "ClienteHTTP heredamos IMPLEMENTACION reutilizable (el COMO enviar requests). De "
    "ITranscriptor heredamos un ROL o CONTRATO (poder usarse de forma polimorfica). Una nos da "
    "el 'como', la otra el 'que'.")
etiqueta("Por que NO herencia simple:",
    "no se puede expresar con herencia simple porque ClienteHTTP e ITranscriptor no se "
    "relacionan: uno es infraestructura de red, el otro un concepto del dominio. Ninguno es "
    "padre del otro. Con herencia simple tendriamos que elegir solo uno y perderiamos o la "
    "reutilizacion de la red, o el polimorfismo.")
etiqueta("El problema del diamante (y como lo evitamos):",
    "la herencia multiple puede causar ambiguedad cuando dos padres aportan el mismo atributo o "
    "metodo (se conoce como 'problema del diamante'). Nosotros lo evitamos siguiendo la buena "
    "practica recomendada en C++: heredar de UNA SOLA clase con implementacion (ClienteHTTP) mas "
    "interfaces PURAS sin estado (ITranscriptor, IObservador). Como las interfaces no tienen "
    "atributos ni logica que choque, no hay diamante.")

# ---------- POLIMORFISMO ----------
h2("2.5 Polimorfismo")
etiqueta("Que es:",
    "la capacidad de tratar objetos de distintas clases a traves de una misma interfaz, "
    "ejecutando el comportamiento correcto de cada uno. Literalmente 'muchas formas'.")
etiqueta("Donde lo usamos:",
    "la GUI guarda un transcriptor y lo usa a traves de los metodos virtuales de la interfaz. "
    "Tambien VentanaPrincipal implementa los metodos de IObservador. La redefinicion se marca "
    "siempre con 'override':")
codigo(
"// ClienteTranscriptor implementa la virtual pura de ITranscriptor\n"
"DIArize::Core::Transcripcion transcribir(\n"
"    const std::string& rutaAudio,\n"
"    const std::string& idioma = \"es\",\n"
"    int numHablantes = 0) override;\n"
"\n"
"// y sobreescribe la virtual con cuerpo\n"
"bool estaDisponible() const override;"
)
etiqueta("Como funciona:",
    "al declarar los metodos como 'virtual', la decision de QUE version ejecutar se toma en "
    "tiempo de EJECUCION segun el tipo real del objeto (esto se llama 'dynamic dispatch'). Si un "
    "puntero ITranscriptor* apunta a un ClienteTranscriptor, al llamar transcribir() se ejecuta "
    "la version de ClienteTranscriptor.")
etiqueta("Por que lo usamos:",
    "para poder cambiar la implementacion sin tocar el codigo que la usa. Manana podriamos crear "
    "un TranscriptorLocal o un TranscriptorMock para tests, y la ventana los usaria igual, "
    "programando contra la interfaz.")
etiqueta("La palabra 'override' (por que importa):",
    "le pide al compilador que VERIFIQUE que realmente estamos sobreescribiendo un metodo "
    "virtual de la base. Si nos equivocamos en la firma (por ejemplo un 'const' de menos), en "
    "vez de crear silenciosamente un metodo nuevo, da un error de compilacion. Es una red de "
    "seguridad.")
etiqueta("Por que NO la otra opcion:",
    "sin polimorfismo, para soportar varios transcriptores tendriamos que llenar el codigo de "
    "'if (tipo == AssemblyAI) ... else if (tipo == Local) ...'. Eso es fragil y hay que tocarlo "
    "cada vez que se agrega un tipo nuevo. El polimorfismo lo resuelve solo.")

# ---------- VIRTUAL PURA VS CON CUERPO ----------
h2("2.6 Funciones virtuales puras vs. con cuerpo")
etiqueta("Que es:",
    "una funcion virtual PURA (termina en '= 0') no tiene implementacion y OBLIGA a las "
    "subclases a darla. Una virtual CON CUERPO tiene un comportamiento por defecto que la "
    "subclase puede o no sobreescribir.")
etiqueta("Donde lo usamos:", "en ITranscriptor conviven las dos:")
codigo(
"virtual Transcripcion transcribir(...) = 0;          // PURA (obligatoria)\n"
"virtual bool estaDisponible() const { return true; } // con cuerpo (opcional)"
)
etiqueta("Como funciona:",
    "transcribir() es pura: todo transcriptor DEBE implementarla, sino la subclase tambien es "
    "abstracta y no se puede instanciar. estaDisponible() tiene un valor por defecto (true) que "
    "ClienteTranscriptor elige sobreescribir para preguntarle al servidor por el endpoint "
    "/estado.")
etiqueta("Por que lo usamos:",
    "para distinguir lo OBLIGATORIO de lo OPCIONAL. Transcribir es la esencia (todo transcriptor "
    "tiene que hacerlo). Estar disponible tiene un default razonable que solo se redefine si hace "
    "falta.")

# ---------- DESTRUCTOR VIRTUAL ----------
h2("2.7 Destructor virtual")
etiqueta("Que es:",
    "un destructor declarado 'virtual' en la clase base, para que al borrar un objeto a traves "
    "de un puntero a la base se ejecute tambien el destructor de la clase hija.")
etiqueta("Donde lo usamos:", "en ITranscriptor, IObservador y ClienteHTTP:")
codigo(
"virtual ~ITranscriptor() = default;\n"
"virtual ~IObservador()   = default;\n"
"virtual ~ClienteHTTP();"
)
etiqueta("Como funciona:",
    "cuando hacemos 'ITranscriptor* p = new ClienteTranscriptor; delete p;' SIN destructor "
    "virtual solo se llamaria el destructor de la base, y todo lo que agrego la hija quedaria sin "
    "liberar (memory leak). CON destructor virtual se llama la cadena completa: primero la hija, "
    "despues la base.")
etiqueta("Por que lo usamos:",
    "es una regla de oro de C++: toda clase base pensada para herencia polimorfica DEBE tener "
    "destructor virtual. De lo contrario, el polimorfismo provoca fugas de memoria.")

# ---------- CLASE ABSTRACTA ----------
h2("2.8 Clase abstracta (no instanciable)")
etiqueta("Que es:",
    "una clase que tiene al menos una funcion virtual pura. El compilador NO permite crear "
    "objetos de ella; solo existe para ser heredada.")
etiqueta("Donde lo usamos:", "ITranscriptor e IObservador son abstractas.")
etiqueta("Como funciona:",
    "intentar 'ITranscriptor x;' da error de compilacion. Lo que si se instancia son sus "
    "implementaciones concretas: ClienteTranscriptor, VentanaPrincipal.")
etiqueta("Por que lo usamos:",
    "porque 'un transcriptor generico' o 'un observador generico' son conceptos, no cosas "
    "concretas. No tiene sentido instanciarlos; lo concreto son sus implementaciones.")

doc.add_page_break()

# ============================================================ PATRONES Y C++ EXTRA
h1("3. Patrones de diseño y herramientas de C++")

h2("3.1 Patron Observer")
etiqueta("Donde:", "la interfaz IObservador (nucleo/include/nucleo/IObservador.h):")
codigo(
"enum class EstadoApp { Inactivo, Subiendo, Procesando, Listo, Error };\n"
"class IObservador {\n"
"public:\n"
"    virtual ~IObservador() = default;\n"
"    virtual void onEstadoCambiado(EstadoApp estado, const std::string& detalle) = 0;\n"
"    virtual void onProgreso(int porcentaje, const std::string& etapa) = 0;\n"
"};"
)
etiqueta("Que es y por que:",
    "el patron Observer desacopla a quien genera eventos de quien reacciona. VentanaPrincipal "
    "implementa onEstadoCambiado y onProgreso, asi se entera de los cambios de estado del "
    "backend sin estar fuertemente atada a el. Permite notificar cambios de forma ordenada.")

h2("3.2 RAII y smart pointers")
etiqueta("Donde:",
    "VentanaPrincipal guarda el transcriptor con un puntero inteligente:")
codigo("std::unique_ptr<DIArize::Red::ClienteTranscriptor> _transcriptor;")
etiqueta("Que es y por que:",
    "RAII significa que la vida de un recurso queda atada a la vida de un objeto. El unique_ptr "
    "libera la memoria AUTOMATICAMENTE cuando se destruye la ventana, sin necesidad de 'delete' "
    "manual. Esto evita fugas de memoria y errores de doble liberacion. La alternativa (punteros "
    "crudos con new/delete a mano) es propensa a olvidos y bugs.")

h2("3.3 const-correctness")
etiqueta("Donde:", "todos los getters del proyecto, por ejemplo getNombre() const.")
etiqueta("Que es y por que:",
    "el 'const' al final de un metodo le promete al compilador que ese metodo NO modifica el "
    "objeto. Si por error intentamos cambiar un atributo adentro, no compila. Da seguridad y "
    "hace el codigo mas legible: quien ve getNombre() const sabe que solo lee.")

h2("3.4 mutable")
etiqueta("Donde:", "el atributo _manager en ClienteHTTP:")
codigo("mutable QNetworkAccessManager _manager;")
etiqueta("Que es y por que:",
    "'mutable' permite modificar un atributo incluso dentro de metodos const. Los metodos get() "
    "y post() son const (logicamente no cambian el cliente), pero QNetworkAccessManager necesita "
    "modificarse internamente para enviar la peticion. mutable resuelve esa contradiccion sin "
    "tener que sacar el const de toda la interfaz.")

h2("3.5 Prohibir la copia con = delete")
etiqueta("Donde:", "en ClienteHTTP:")
codigo(
"ClienteHTTP(const ClienteHTTP&)            = delete;\n"
"ClienteHTTP& operator=(const ClienteHTTP&) = delete;"
)
etiqueta("Que es y por que:",
    "'= delete' prohibe explicitamente copiar el objeto: si alguien lo intenta, da error de "
    "compilacion. Lo usamos porque QNetworkAccessManager no es copiable. Es mejor un error claro "
    "en compilacion que un bug raro en ejecucion. (Aclaracion: esto NO es sobrecarga de "
    "operadores; al contrario, estamos eliminando el operador de copia.)")

h2("3.6 Contenedores de la STL")
parrafo("Elegimos cada estructura segun como se usa:")
tabla2 = doc.add_table(rows=1, cols=3)
tabla2.style = "Light Grid Accent 1"
h = tabla2.rows[0].cells
h[0].paragraphs[0].add_run("Estructura").bold = True
h[1].paragraphs[0].add_run("Donde").bold = True
h[2].paragraphs[0].add_run("Por que esa y no otra").bold = True
filas2 = [
    ("std::vector<Segmento>", "Transcripcion",
     "Los segmentos se recorren en orden y se accede por indice. El vector guarda los datos en "
     "memoria contigua, ideal para iterar rapido."),
    ("std::map<string,double>", "participacion por hablante",
     "Asocia nombre de hablante -> segundos. Es clave-valor con busqueda por clave."),
    ("std::list<Transcripcion>", "Repositorio (historial)",
     "Insercion O(1) al final y solo se recorre secuencialmente; no necesitamos acceso por "
     "indice, asi que la lista enlazada es lo adecuado."),
]
for a, b, c in filas2:
    row = tabla2.add_row().cells
    row[0].paragraphs[0].add_run(a)
    row[1].paragraphs[0].add_run(b)
    row[2].paragraphs[0].add_run(c)
doc.add_paragraph()
etiqueta("vector vs list (si preguntan):",
    "vector = memoria contigua, acceso rapido por indice [i], pero insertar en el medio es caro. "
    "list = lista enlazada, insertar/borrar en cualquier lado es O(1), pero no permite [i]. "
    "Elegimos cada una segun el uso real.")

doc.add_page_break()

# ============================================================ CONEXION API
h1("4. Como nos conectamos a las APIs")
parrafo("Hay dos APIs externas y un servidor intermedio. La cadena completa es:")
codigo(
"GUI C++    --HTTP-->    Servidor Flask (Python)    --HTTPS-->    AssemblyAI / Gemini\n"
"(cliente)              (localhost:8765)                          (la nube)"
)
h2("4.1 Por que un servidor Python en el medio (y no llamar directo desde C++)")
bullet("Seguridad de las API keys: las claves viven en el .env del servidor, nunca en el "
       "ejecutable C++. Si distribuimos el .exe, nadie roba las claves.")
bullet("Reutilizacion de librerias: AssemblyAI y OpenAI tienen SDKs oficiales maduros en "
       "Python; en C++ habria que armar todo a mano.")
bullet("Separacion de responsabilidades: C++ se encarga de la interfaz; Python de la IA y el "
       "procesamiento.")

h2("4.2 Conexion HTTP REST (modo archivo)")
parrafo(
    "ClienteHTTP encapsula QNetworkAccessManager de Qt y ofrece get(), post() (con "
    "QHttpMultiPart para subir el archivo de audio) y postJson() (para las operaciones de IA). "
    "Qt es asincrono por naturaleza, pero nosotros queriamos llamadas sincronas simples; lo "
    "resolvimos con un QEventLoop que bloquea hasta que llega la respuesta:")
codigo(
"QEventLoop loop;\n"
"QNetworkReply* reply = _manager.get(construirRequest(url));\n"
"QObject::connect(reply, &QNetworkReply::finished, &loop, &QEventLoop::quit);\n"
"loop.exec();   // bloquea aca hasta que termina el request"
)
parrafo(
    "El servidor recibe el audio en el endpoint /transcribir, lo manda a AssemblyAI mediante el "
    "PipelineAssemblyAI y devuelve un JSON con texto, segmentos y participacion. El C++ parsea "
    "ese JSON con QJsonDocument.")

h2("4.3 Conexion WebSocket (modo en vivo)")
parrafo(
    "Para la transcripcion en vivo no sirve el esquema pedido-respuesta de HTTP: hace falta un "
    "canal bidireccional y continuo. Por eso usamos WebSocket (QWebSocket) en ClienteStreaming. "
    "Capturamos el microfono con QAudioSource (PCM 16-bit mono 16kHz, el formato que pide "
    "AssemblyAI), enviamos chunks de ~100ms y recibimos mensajes con texto parcial o final.")
etiqueta("El truco del token efimero:",
    "el C++ NO se conecta al WebSocket con la API key. Primero le pide al servidor un token "
    "temporal (endpoint /streaming-token) que dura 10 minutos. Asi la clave real nunca sale del "
    "servidor y el cliente streamea de forma segura.")

doc.add_page_break()

# ============================================================ RESUMEN
h1("5. Resumen rapido (cheat-sheet)")
parrafo("Conceptos de POO aplicados, con un ejemplo de cada uno:")
resumen = [
    ("Encapsulamiento", "atributos private + getters/setters (Hablante, Segmento)."),
    ("Abstraccion", "clases abstractas ITranscriptor e IObservador."),
    ("Herencia simple", "ClienteTranscriptor hereda de ClienteHTTP (reutiliza la red)."),
    ("Herencia multiple", "ClienteTranscriptor (HTTP + interfaz) y VentanaPrincipal (ventana + observador)."),
    ("Polimorfismo", "metodos virtual + override, dispatch dinamico via ITranscriptor."),
    ("Virtual pura vs con cuerpo", "transcribir() = 0 (obligatoria) vs estaDisponible() (default)."),
    ("Destructor virtual", "en las clases base, evita fugas de memoria."),
    ("Clase abstracta", "ITranscriptor / IObservador no se pueden instanciar."),
    ("RAII / smart pointers", "std::unique_ptr para el transcriptor."),
    ("const-correctness", "todos los getters son const."),
    ("mutable", "_manager en ClienteHTTP."),
    ("= delete", "prohibir copia en ClienteHTTP."),
    ("Patron Observer", "IObservador para notificar cambios de estado."),
    ("STL", "vector, list, map, string, pair."),
]
for nombre, desc in resumen:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(nombre + ": ")
    r.font.bold = True
    r.font.color.rgb = CYAN
    p.add_run(desc)

doc.add_paragraph()
parrafo("Lo que NO usamos (por si preguntan):", bold=True)
bullet("Sobrecarga de operadores: no se uso. Nuestras clases se manejan con metodos con nombre "
       "(agregarSegmento, getTexto), mas legibles que un operador en este dominio.")
bullet("Templates propios: usamos los de la STL (vector, map), pero no definimos clases template "
       "propias.")
bullet("Herencia virtual: no hizo falta porque evitamos el problema del diamante por diseño.")

ruta = r"c:/Users/octav/OneDrive/Escritorio/ProyectoFinalPoo/DIArize_POO_Brachetti_Diaz_Lopardo/Conceptos_POO_DIArize.docx"
doc.save(ruta)
print("Documento generado en:", ruta)
